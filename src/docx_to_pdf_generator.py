# Resume_Tailoring/src/docx_to_pdf_generator.py
import logging
import re
from typing import Dict, List, Any, Optional
import os
import tempfile # Keep for potential DOCX generation before upload
import json
# --- Google Drive API Imports ---
import io
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from google.auth.transport.requests import Request as AuthRequest # Added

# --- Document Generation Imports (python-docx) ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

# --- Configuration Import ---
app_config = None # Initialize
try:
    import config as app_config # Absolute import from project root
    # SERVICE_ACCOUNT_FILE_PATH is not directly used here anymore for the flag,
    # but app_config itself is crucial for SERVICE_ACCOUNT_JSON_CONTENT.
    logging.info("docx_to_pdf_generator: Successfully imported app_config.")
except ImportError as e:
    logging.warning(f"docx_to_pdf_generator: Could not import app_config via absolute import: {e}. Google Drive features will likely fail.")
    # app_config remains None

logger = logging.getLogger(__name__)
EMU_PER_POINT = Pt(1)

# --- Google Drive API Helper Functions (from your example) ---
SCOPES = ['https://www.googleapis.com/auth/drive']

def get_drive_service():
    """Authenticates and returns a Google Drive API service object."""
    creds = None
    
    # Access the JSON content string from app_config
    service_account_json_content_str = getattr(app_config, 'SERVICE_ACCOUNT_JSON_CONTENT', None)

    if not service_account_json_content_str or not isinstance(service_account_json_content_str, str):
        logger.error("Service account JSON content is not set or not a string in config.")
        logger.error("Please set the SERVICE_ACCOUNT_JSON_CONTENT variable in config.py correctly.")
        return None
        
    try:
        # Parse the JSON string into a dictionary
        service_account_info = json.loads(service_account_json_content_str)
        
        # Create credentials from the parsed info
        creds = service_account.Credentials.from_service_account_info(
            service_account_info, scopes=SCOPES
        )
        logger.info("Successfully created Google Drive credentials from JSON content.")
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse SERVICE_ACCOUNT_JSON_CONTENT: {e}", exc_info=True)
        return None
    except Exception as e:
        logger.error(f"Error creating credentials from service account info: {e}", exc_info=True)
        return None

    try:
        service = build('drive', 'v3', credentials=creds, static_discovery=False)
        logger.info("Google Drive API service created successfully.")
        return service
    except Exception as e:
        logger.error(f"Failed to build Google Drive service: {e}", exc_info=True)
        return None

def upload_and_convert_to_google_doc(drive_service, local_docx_path, drive_filename_prefix):
    """
    Uploads a local DOCX file to Google Drive, then copies it to create a native Google Doc.
    Returns the File ID of the native Google Doc and the ID of the original uploaded DOCX.
    """
    if not os.path.exists(local_docx_path):
        logger.error(f"Local DOCX file not found: {local_docx_path}")
        return None, None

    original_docx_file_metadata = {
        'name': f"{drive_filename_prefix}_original.docx",
        'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    media = MediaFileUpload(local_docx_path,
                            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            resumable=True)
    original_uploaded_file_id = None
    try:
        logger.info(f"Uploading original DOCX '{original_docx_file_metadata['name']}' to Google Drive...")
        original_file = drive_service.files().create(
            body=original_docx_file_metadata,
            media_body=media,
            fields='id, name'
        ).execute()
        original_uploaded_file_id = original_file.get('id')
        logger.info(f"Original DOCX uploaded. Drive File ID: {original_uploaded_file_id}, Name: {original_file.get('name')}")
        if not original_uploaded_file_id:
            logger.error("Failed to upload original DOCX or get its ID.")
            return None, None
    except Exception as e:
        logger.error(f"Error uploading original DOCX to Google Drive: {e}", exc_info=True)
        return None, None

    google_doc_native_id = None
    try:
        copy_metadata = {
            'name': f"{drive_filename_prefix}_native_gdoc",
            'mimeType': 'application/vnd.google-apps.document'
        }
        logger.info(f"Converting/Copying uploaded DOCX (ID: {original_uploaded_file_id}) to a native Google Doc format...")
        copied_file = drive_service.files().copy(
            fileId=original_uploaded_file_id,
            body=copy_metadata,
            fields='id, name'
        ).execute()
        google_doc_native_id = copied_file.get('id')
        logger.info(f"Successfully converted to native Google Doc. New File ID: {google_doc_native_id}, Name: {copied_file.get('name')}")
        return original_uploaded_file_id, google_doc_native_id
    except Exception as e:
        logger.error(f"Error converting/copying DOCX to Google Doc format: {e}", exc_info=True)
        return original_uploaded_file_id, None

def export_pdf_from_drive(drive_service, file_id, local_pdf_path):
    """Exports a file from Google Drive as PDF and saves it locally."""
    try:
        logger.info(f"Exporting Drive File ID '{file_id}' as PDF to '{local_pdf_path}'...")
        request = drive_service.files().export_media(fileId=file_id, mimeType='application/pdf')
        output_dir = os.path.dirname(local_pdf_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            logger.info(f"Created output directory: {output_dir}")

        fh = io.FileIO(local_pdf_path, mode='wb')
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            if status:
                logger.info(f"PDF Download progress: {int(status.progress() * 100)}%")
        fh.close()
        if os.path.exists(local_pdf_path) and os.path.getsize(local_pdf_path) > 0:
            logger.info(f"PDF exported and saved successfully to: {os.path.abspath(local_pdf_path)}")
            return True
        else:
            logger.error(f"PDF export failed: Output file '{local_pdf_path}' not found or is empty after download attempt.")
            return False
    except Exception as e:
        logger.error(f"Error exporting PDF from Google Drive: {e}", exc_info=True)
        return False

def delete_file_from_drive(drive_service, file_id):
    """Deletes a file from Google Drive."""
    if not file_id:
        logger.warning("No file_id provided for deletion.")
        return
    try:
        logger.info(f"Deleting Drive File ID '{file_id}'...")
        drive_service.files().delete(fileId=file_id).execute()
        logger.info(f"File deleted successfully from Google Drive (ID: {file_id}).")
    except Exception as e:
        logger.warning(f"Could not delete file from Google Drive (ID: {file_id}): {e}", exc_info=True)


# --- Existing Helper functions (add_hyperlink, add_runs_with_markdown_bold, add_styled_paragraph) ---
# Keep these as they are, they are used for DOCX generation.
def add_hyperlink(paragraph, url, text, font_name='Times New Roman', font_size=Pt(10), color_hex=None, is_bold=False, is_underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    if font_size:
        point_value = int(font_size / EMU_PER_POINT)
        size_el = OxmlElement('w:sz')
        size_el.set(qn('w:val'), str(point_value * 2))
        rPr.append(size_el)
        size_cs_el = OxmlElement('w:szCs')
        size_cs_el.set(qn('w:val'), str(point_value * 2))
        rPr.append(size_cs_el)
    if color_hex:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), color_hex.replace("#",""))
        rPr.append(color_el)
    if is_bold:
        bold_el = OxmlElement('w:b')
        rPr.append(bold_el)
    if is_underline:
        underline_el = OxmlElement('w:u')
        underline_el.set(qn('w:val'), 'single')
        rPr.append(underline_el)
    new_run_el.append(rPr)
    new_run_el.text = text
    hyperlink.append(new_run_el)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_runs_with_markdown_bold(paragraph, text_with_markdown: str,
                                font_name: Optional[str],
                                font_size: Optional[Pt],
                                base_bold: bool = False,
                                base_italic: bool = False):
    parts = re.split(r'(\*\*[^*]+\*\*)', text_with_markdown)
    for part_idx, part_text in enumerate(parts):
        if not part_text:
            continue
        run = paragraph.add_run()
        if font_name:
            run.font.name = font_name
            rpr = run._r.get_or_add_rPr()
            rFonts_el = rpr.get_or_add_rFonts()
            for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')]:
                if attr in rFonts_el.attrib: del rFonts_el.attrib[attr]
            rFonts_el.set(qn('w:ascii'), font_name)
            rFonts_el.set(qn('w:hAnsi'), font_name)
            rFonts_el.set(qn('w:cs'), font_name)
            rFonts_el.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = font_size
        run.italic = base_italic
        text_to_add = ""
        if part_text.startswith('**') and part_text.endswith('**'):
            text_to_add = part_text[2:-2]
            run.bold = True
        else:
            text_to_add = part_text
            run.bold = base_bold
        final_text = text_to_add.replace('**', '')
        run.text = final_text

def add_styled_paragraph(document, text: str, style_name: Optional[str] = None,
                         font_name: str = 'Times New Roman',
                         font_size: Optional[Pt] = None,
                         is_bold: Optional[bool] = None,
                         is_italic: Optional[bool] = None,
                         alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
                         space_after: Optional[Pt] = None, space_before: Optional[Pt] = None,
                         line_spacing: Optional[float] = None, keep_with_next: Optional[bool] = None):
    p = document.add_paragraph(style=style_name)
    if alignment is not None: p.alignment = alignment
    if space_before is not None: p.paragraph_format.space_before = space_before
    if space_after is not None: p.paragraph_format.space_after = space_after
    if line_spacing is not None: p.paragraph_format.line_spacing = line_spacing
    if keep_with_next is not None: p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.widow_control = True
    current_base_bold = is_bold if is_bold is not None else False
    current_base_italic = is_italic if is_italic is not None else False
    add_runs_with_markdown_bold(p, text, font_name, font_size,
                                base_bold=current_base_bold,
                                base_italic=current_base_italic)
    return p

# --- Existing DOCX Section Adding Functions ---
# (add_contact_info_docx, add_section_header_docx, add_summary_docx, etc.)
# These should largely remain the same, as they format the DOCX content.
def add_contact_info_docx(document, contact_data: Dict[str, str]):
    logger.info("Adding contact information to DOCX...")
    add_styled_paragraph(document, contact_data.get("name", "Candidate Name"),
                         font_name='Times New Roman', font_size=Pt(16), is_bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(3))

    line1_full_text = contact_data.get("line1_info", "")
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    match = re.search(email_pattern, line1_full_text)
    p_line1 = document.add_paragraph()
    p_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line1.paragraph_format.space_after = Pt(1)
    p_line1.paragraph_format.widow_control = True
    current_font_name = 'Times New Roman'
    current_font_size = Pt(10)

    if match:
        email_address = match.group(0)
        parts = line1_full_text.split(email_address, 1)
        if parts[0]:
            run_before = p_line1.add_run(parts[0])
            run_before.font.name = current_font_name; run_before.font.size = current_font_size
        add_hyperlink(p_line1, f"mailto:{email_address}", email_address,
                      font_name=current_font_name, font_size=current_font_size,
                      color_hex="0563C1", is_underline=True)
        if parts[1]:
            run_after = p_line1.add_run(parts[1])
            run_after.font.name = current_font_name; run_after.font.size = current_font_size
    else:
        run_line1_plain = p_line1.add_run(line1_full_text)
        run_line1_plain.font.name = current_font_name; run_line1_plain.font.size = current_font_size

    p_links = document.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_after = Pt(18)
    p_links.paragraph_format.widow_control = True
    add_hyperlink(p_links, contact_data.get("linkedin_url", "#"), contact_data.get("linkedin_text", "linkedin.com"),
                  font_name=current_font_name, font_size=current_font_size, color_hex="0563C1", is_underline=True)
    run_sep1 = p_links.add_run(" | "); run_sep1.font.name = current_font_name; run_sep1.font.size = current_font_size
    add_hyperlink(p_links, contact_data.get("github_url", "#"), contact_data.get("github_text", "GitHub"),
                  font_name=current_font_name, font_size=current_font_size, color_hex="0563C1", is_underline=True)
    run_sep2 = p_links.add_run(" | "); run_sep2.font.name = current_font_name; run_sep2.font.size = current_font_size
    add_hyperlink(p_links, contact_data.get("portfolio_url", "#"), contact_data.get("portfolio_text", "Portfolio"),
                  font_name=current_font_name, font_size=current_font_size, color_hex="0563C1", is_underline=True)

def add_section_header_docx(document, header_text: str):
    p = add_styled_paragraph(document, header_text.upper(), font_name='Times New Roman',
                             font_size=Pt(10), is_bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(6), space_after=Pt(4))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single'); bottom_border.set(qn('w:sz'), '4')
    bottom_border.set(qn('w:space'), '1'); bottom_border.set(qn('w:color'), '444444')
    pBdr.append(bottom_border); pPr.append(pBdr)

def add_summary_docx(document, summary_text: str):
    logger.info("Adding summary to DOCX...")
    add_section_header_docx(document, "SUMMARY")
    add_styled_paragraph(document, summary_text, font_name='Times New Roman', font_size=Pt(10),
                         line_spacing=1.15, space_after=Pt(6))

def add_work_experience_docx(document, work_experience_text: str):
    logger.info("Adding work experience to DOCX...")
    add_section_header_docx(document, "WORK EXPERIENCE")
    if not work_experience_text or not work_experience_text.strip():
        logger.warning("Work experience text is empty. Skipping section content.")
        add_styled_paragraph(document, "N/A", font_size=Pt(10))
        return

    cleaned_text = re.sub(r"^## Work Experience\s*\n+", "", work_experience_text, flags=re.IGNORECASE).strip()
    job_entries = re.split(r'\n\s*\n+(?=\s*(?:\*\*)?[A-Z][\w\s.,-]+?\s*(?:\*\*)?\s*\|)', cleaned_text)
    right_tab_stop = Inches(7.4) # Adjust as needed

    for entry_text in job_entries:
        entry_text = entry_text.strip()
        if not entry_text: continue
        lines = [line.strip() for line in entry_text.split('\n') if line.strip()]
        if not lines: continue

        p_job_header = document.add_paragraph()
        p_job_header.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_ALIGN_PARAGRAPH.RIGHT)
        p_job_header.paragraph_format.widow_control = True
        header_parts = lines[0].split('|')
        add_runs_with_markdown_bold(p_job_header, header_parts[0].strip(), 'Times New Roman', Pt(10), base_bold=True)
        if len(header_parts) > 1:
            p_job_header.add_run(" | ").font.name = 'Times New Roman';
            add_runs_with_markdown_bold(p_job_header, header_parts[1].strip(), 'Times New Roman', Pt(10), base_bold=True)
        if len(header_parts) > 2:
            p_job_header.add_run(" | ").font.name = 'Times New Roman';
            add_runs_with_markdown_bold(p_job_header, header_parts[2].strip(), 'Times New Roman', Pt(10), base_italic=True)

        bullet_start_index = 1
        if len(lines) > 1 and re.match(r'^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|September)[\w\s,–-]+$', lines[1].strip(), re.IGNORECASE):
            date_line_str = lines[1].strip()
            run_date = p_job_header.add_run('\t' + date_line_str)
            run_date.font.name = 'Times New Roman'; run_date.font.size = Pt(10); run_date.italic = True
            bullet_start_index = 2
        
        p_job_header.paragraph_format.space_after = Pt(2)
        p_job_header.paragraph_format.keep_with_next = True

        for i in range(bullet_start_index, len(lines)):
            line = lines[i]
            if line.startswith('*'):
                add_styled_paragraph(document, line[1:].strip(), style_name='List Bullet',
                                     font_name='Times New Roman', font_size=Pt(10),
                                     space_after=Pt(2), line_spacing=1.15)
        if len(lines) > bullet_start_index and document.paragraphs:
            document.paragraphs[-1].paragraph_format.space_after = Pt(6)


def add_technical_skills_docx(document, skills_text: str):
    logger.info("Adding technical skills to DOCX...")
    add_section_header_docx(document, "TECHNICAL SKILLS")
    if not skills_text or not skills_text.strip():
        logger.warning("Technical skills text is empty. Skipping section content.")
        add_styled_paragraph(document, "N/A", font_size=Pt(10))
        return
        
    cleaned_text = re.sub(r"^## Technical Skills\s*\n+", "", skills_text, flags=re.IGNORECASE).strip()
    categories = cleaned_text.split('\n')
    for category_line in categories:
        category_line = category_line.strip()
        if not category_line: continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.widow_control = True
        parts = category_line.split(':', 1)
        if len(parts) == 2:
            add_runs_with_markdown_bold(p, parts[0].strip() + ": ", 'Times New Roman', Pt(10), base_bold=True)
            add_runs_with_markdown_bold(p, parts[1].strip(), 'Times New Roman', Pt(10))
        else:
            add_runs_with_markdown_bold(p, category_line, 'Times New Roman', Pt(10))

def add_projects_docx(document, projects_text: str, contact_data: Dict[str, str]):
    logger.info("Adding projects to DOCX...")
    add_section_header_docx(document, "PROJECTS")
    if not projects_text or not projects_text.strip():
        logger.warning("Projects text is empty. Skipping section content.")
        add_styled_paragraph(document, "N/A", font_size=Pt(10))
        return

    cleaned_projects_text = re.sub(r"^(## |### )[\w\s]*Projects?[\w\s]*(?:Section)?:?\s*\n*", "", projects_text.strip(), flags=re.IGNORECASE).strip()
    if not cleaned_projects_text:
        logger.warning("Projects text is empty after cleaning headers.")
        return

    github_base_url = contact_data.get("github_url", "https://github.com/DefaultUser")
    project_hyperlinks = {
        "Intelligent Building Code QA": "https://virginia-building-codes.streamlit.app/",
        "AI-Text Discriminator": f"{github_base_url}/AI-Content-Filter"
    }
    
    project_entries = re.split(r'\n\s*\n+(?=\s*(?:\*\*)?[A-Z0-9][\w\s\-()]+?(?:\*\*)?\s*(?:\||\n))', cleaned_projects_text)
    
    for entry_text in project_entries:
        entry_text = entry_text.strip()
        if not entry_text: continue
        lines = [line.strip() for line in entry_text.split('\n') if line.strip()]
        if not lines: continue
            
        title_line = lines[0]
        title_parts = title_line.split('|', 1)
        project_name_raw = re.sub(r"^\*\*(.*?)\*\*$", r"\1", title_parts[0].strip())

        p_title = document.add_paragraph()
        p_title.paragraph_format.widow_control = True
        project_url = project_hyperlinks.get(project_name_raw)
        
        if project_url:
            add_hyperlink(p_title, project_url, project_name_raw,
                          font_name='Times New Roman', font_size=Pt(10),
                          is_bold=True, color_hex="0563C1", is_underline=True)
        else:
            add_runs_with_markdown_bold(p_title, project_name_raw, 'Times New Roman', Pt(10), base_bold=True)
            logger.warning(f"No hyperlink found for project: {project_name_raw}")
            
        if len(title_parts) > 1 and title_parts[1].strip():
            tagline_text = title_parts[1].strip().replace("_", "")
            p_title.add_run(" | ").font.name = 'Times New Roman'
            add_runs_with_markdown_bold(p_title, tagline_text, 'Times New Roman', Pt(10), base_italic=True)
            
        p_title.paragraph_format.space_after = Pt(2)
        p_title.paragraph_format.keep_with_next = True
        
        for i in range(1, len(lines)):
            line = lines[i]
            if line.startswith('*'):
                add_styled_paragraph(document, line[1:].strip(), style_name='List Bullet',
                                     font_name='Times New Roman', font_size=Pt(10),
                                     space_after=Pt(2), line_spacing=1.15)
        if len(lines) > 1 and document.paragraphs:
            if document.paragraphs[-1].text.strip() == lines[-1][1:].strip().replace("**",""):
                document.paragraphs[-1].paragraph_format.space_after = Pt(6)

def add_education_docx(document, education_list: List[Dict[str, str]]):
    logger.info("Adding education to DOCX...")
    add_section_header_docx(document, "EDUCATION")
    if not education_list:
        logger.warning("Education list is empty. Skipping section content.")
        add_styled_paragraph(document, "N/A", font_size=Pt(10))
        return

    right_tab_stop = Inches(7.4) # Adjust as needed
    for edu_item in education_list:
        p_edu_line = document.add_paragraph()
        p_edu_line.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_ALIGN_PARAGRAPH.RIGHT)
        p_edu_line.paragraph_format.widow_control = True
        degree_str = edu_item.get('degree_line', 'Degree N/A')
        uni_str = edu_item.get('university_line', 'University N/A')
        dates_str = edu_item.get("dates_line", "Dates N/A")
        add_runs_with_markdown_bold(p_edu_line, degree_str, 'Times New Roman', Pt(10), base_bold=True)
        p_edu_line.add_run(", ").font.name = 'Times New Roman'
        add_runs_with_markdown_bold(p_edu_line, uni_str, 'Times New Roman', Pt(10))
        run_dates = p_edu_line.add_run('\t' + dates_str)
        run_dates.font.name = 'Times New Roman'; run_dates.font.size = Pt(10); run_dates.italic = True
        p_edu_line.paragraph_format.space_after = Pt(6)


# --- Modified PDF Generation Functions to use Google Drive ---

def generate_pdf_via_google_drive(
    document: Document, # The python-docx Document object
    output_pdf_directory: str,
    base_filename: str # e.g., "Venkatesh_Shanmugam_Resume_TargetCompany_AI_4YOE"
) -> Optional[str]:
    """
    Saves a DOCX Document to a temporary file, uploads to Google Drive,
    converts to Google Doc, exports as PDF, saves it locally, and cleans up Drive files.
    Returns the path to the locally saved PDF or None on failure.
    """
    drive_service = get_drive_service()
    if not drive_service:
        logger.error("Could not get Google Drive service. PDF generation via Drive failed.")
        return None

    temp_docx_file_path = None
    original_docx_id_on_drive = None
    native_gdoc_id_on_drive = None
    final_pdf_filepath = None

    try:
        # 1. Save the python-docx Document to a temporary local DOCX file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="temp_resume_") as tmpfile:
            temp_docx_file_path = tmpfile.name
        document.save(temp_docx_file_path)
        logger.info(f"Temporary DOCX for Drive conversion saved to: {temp_docx_file_path}")

        # 2. Upload DOCX and convert to Google Doc
        # Use a prefix derived from base_filename for Drive files
        drive_filename_prefix = re.sub(r'[^\w\-_]', '_', base_filename) # Sanitize
        original_docx_id_on_drive, native_gdoc_id_on_drive = upload_and_convert_to_google_doc(
            drive_service, temp_docx_file_path, drive_filename_prefix
        )

        if not native_gdoc_id_on_drive:
            logger.error("Failed to create a native Google Doc from the uploaded DOCX. PDF generation via Drive failed.")
            return None # Critical failure

        # 3. Export the native Google Doc as PDF
        # Ensure output directory exists
        os.makedirs(output_pdf_directory, exist_ok=True)
        final_pdf_filepath = os.path.join(output_pdf_directory, f"{base_filename}.pdf")

        pdf_export_successful = export_pdf_from_drive(
            drive_service, native_gdoc_id_on_drive, final_pdf_filepath
        )

        if pdf_export_successful:
            logger.info(f"PDF generated successfully via Google Drive: '{final_pdf_filepath}'")
            return final_pdf_filepath
        else:
            logger.error("PDF export from native Google Doc failed.")
            return None

    except Exception as e:
        logger.error(f"Error during Google Drive PDF generation pipeline: {e}", exc_info=True)
        return None
    finally:
        # 4. Clean up: Delete files from Drive
        if drive_service: # Only attempt deletion if service was obtained
            if native_gdoc_id_on_drive:
                delete_file_from_drive(drive_service, native_gdoc_id_on_drive)
            if original_docx_id_on_drive:
                delete_file_from_drive(drive_service, original_docx_id_on_drive)
        
        # 5. Clean up: Delete local temporary DOCX file
        if temp_docx_file_path and os.path.exists(temp_docx_file_path):
            try:
                os.remove(temp_docx_file_path)
                logger.info(f"Local temporary DOCX '{temp_docx_file_path}' removed.")
            except OSError as e_remove:
                logger.warning(f"Could not remove local temporary DOCX: {e_remove}")


# --- Main PDF Generation Functions (MODIFIED to use Google Drive conversion) ---

def generate_styled_resume_pdf(
    tailored_data: Dict[str, Any],
    contact_info: Dict[str, str],
    education_info: List[Dict[str, str]],
    output_pdf_directory: str,
    target_company_name: Optional[str] = None, # Used in filename
    years_of_experience: Optional[int] = None, # Used in filename
    filename_keyword: str = "Resume" # Base keyword for the filename
) -> Optional[str]:

    logger.info(f"Starting styled RESUME PDF generation using Google Drive. Output dir: '{output_pdf_directory}'")
    document = Document()
    # (Setup styles and margins as before - this part is for python-docx DOCX creation)
    normal_style = document.styles['Normal']
    normal_font = normal_style.font; normal_font.name = 'Times New Roman'; normal_font.size = Pt(10)
    ct_style_rpr = normal_style.element.get_or_add_rPr()
    ct_style_fonts = ct_style_rpr.get_or_add_rFonts()
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        if attr in ct_style_fonts.attrib: del ct_style_fonts.attrib[attr]
    ct_style_fonts.set(qn('w:ascii'), 'Times New Roman'); ct_style_fonts.set(qn('w:hAnsi'), 'Times New Roman')
    ct_style_fonts.set(qn('w:cs'), 'Times New Roman'); ct_style_fonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal_style.paragraph_format.space_before = Pt(0); normal_style.paragraph_format.space_after = Pt(2)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    try:
        list_bullet_style = document.styles['List Bullet']
        lb_font = list_bullet_style.font; lb_font.name = 'Times New Roman'; lb_font.size = Pt(10)
        lb_ct_style_rpr = list_bullet_style.element.get_or_add_rPr()
        lb_ct_style_fonts = lb_ct_style_rpr.get_or_add_rFonts()
        for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')]:
            if attr in lb_ct_style_fonts.attrib: del lb_ct_style_fonts.attrib[attr]
        lb_ct_style_fonts.set(qn('w:ascii'), 'Times New Roman'); lb_ct_style_fonts.set(qn('w:hAnsi'), 'Times New Roman')
        lb_ct_style_fonts.set(qn('w:cs'), 'Times New Roman'); lb_ct_style_fonts.set(qn('w:eastAsia'), 'Times New Roman')
        list_bullet_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        list_bullet_style.paragraph_format.line_spacing = 1.15
        list_bullet_style.paragraph_format.left_indent = Inches(0.25)
        list_bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    except KeyError:
        logger.warning("'List Bullet' style not found. Bulleted lists may not be properly formatted.")

    for section_elm in document.sections:
        section_elm.left_margin = Inches(0.51); section_elm.right_margin = Inches(0.51)
        section_elm.top_margin = Inches(0.13); section_elm.bottom_margin = Inches(0.06)
    
    # Add content to the DOCX document
    add_contact_info_docx(document, contact_info)
    if tailored_data.get("summary"): add_summary_docx(document, tailored_data["summary"])
    if tailored_data.get("work_experience"): add_work_experience_docx(document, tailored_data["work_experience"])
    if tailored_data.get("technical_skills"): add_technical_skills_docx(document, tailored_data["technical_skills"])
    if tailored_data.get("projects"): add_projects_docx(document, tailored_data["projects"], contact_info)
    add_education_docx(document, education_info)

    # Construct the base filename for the PDF
    candidate_last_name = contact_info.get("name", "Candidate").split()[-1] if contact_info.get("name") else "Resume"
    yoe_str = str(years_of_experience) if years_of_experience is not None else "X"
    company_str = re.sub(r'\W+', '', target_company_name) if target_company_name else "TargetCompany"
    
    # filename_keyword can be a role, e.g., "AI_Engineer"
    # The full filename_keyword passed from scrape.py might already be very descriptive.
    # If filename_keyword already contains company/role, this might be redundant.
    # Let's assume filename_keyword is a general keyword like "AI" or "DataScience" for now if scrape.py isn't updated.
    # If scrape.py sends a full prefix in filename_keyword, this becomes simpler.
    # For now, using the logic from your original main.py for filename parts.
    # If filename_keyword already contains company/role, this might be redundant.
    # Let's assume filename_keyword is a general keyword like "AI" or "DataScience" for now if scrape.py isn't updated.
    # If scrape.py sends a full prefix in filename_keyword, this becomes simpler.
    # For now, using the logic from your original main.py for filename parts.
    base_pdf_filename = f"{filename_keyword}_{company_str}_{candidate_last_name}_{yoe_str}YOE"
    base_pdf_filename = re.sub(r'[^\w\.\-_]', '_', base_pdf_filename) # Sanitize

    return generate_pdf_via_google_drive(document, output_pdf_directory, base_pdf_filename)

# In Resume_Tailoring/src/docx_to_pdf_generator.py
# Make sure all necessary imports like re, os, logging, Document, Pt, Inches, WD_ALIGN_PARAGRAPH, qn,
# add_styled_paragraph, add_hyperlink, generate_pdf_via_google_drive are present at the top of your file.

logger = logging.getLogger(__name__) # Ensure logger is defined

def generate_cover_letter_pdf(
    cover_letter_body_text: str,
    contact_info: Dict[str, str],
    job_title: str,
    company_name: str,
    output_pdf_directory: str,
    filename_keyword: str = "CoverLetter",
    years_of_experience: Optional[int] = None
) -> Optional[str]:

    logger.info(f"Starting styled COVER LETTER PDF generation. Output dir: '{output_pdf_directory}'")
    document = Document()

    # --- Default Styles and Margins for Cover Letter ---
    # (Your existing style setup from the previous corrected version, ensuring normal_style.paragraph_format.widow_control = True)
    normal_style = document.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(11)

    ct_style_rpr = normal_style.element.get_or_add_rPr()
    ct_style_fonts = ct_style_rpr.get_or_add_rFonts()
    theme_font_attributes = [
        qn('w:asciiTheme'), qn('w:hAnsiTheme'),
        qn('w:eastAsiaTheme'), qn('w:cstheme')
    ]
    for attr in theme_font_attributes:
        if attr in ct_style_fonts.attrib:
            del ct_style_fonts.attrib[attr]
    
    ct_style_fonts.set(qn('w:ascii'), 'Times New Roman')
    ct_style_fonts.set(qn('w:hAnsi'), 'Times New Roman')
    ct_style_fonts.set(qn('w:cs'), 'Times New Roman')
    ct_style_fonts.set(qn('w:eastAsia'), 'Times New Roman')

    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal_style.paragraph_format.widow_control = True

    for section_elm in document.sections:
        section_elm.left_margin = Inches(1)
        section_elm.right_margin = Inches(1)
        section_elm.top_margin = Inches(0.75) 
        section_elm.bottom_margin = Inches(0.75)

    # --- Create Traditional Letterhead Style Top Contact Block (Street Address Omitted) ---
    letterhead_font_name = 'Times New Roman'
    letterhead_font_size = Pt(11) 
    letterhead_name_font_size = Pt(14)
    letterhead_alignment = WD_ALIGN_PARAGRAPH.LEFT # Or .CENTER or .RIGHT as you prefer

    # Candidate Name
    if contact_info.get("name"):
        add_styled_paragraph( # Assuming add_styled_paragraph sets widow_control
            document,
            contact_info["name"],
            font_name=letterhead_font_name,
            font_size=letterhead_name_font_size,
            is_bold=True,
            alignment=letterhead_alignment,
            space_after=Pt(1) 
        )

    # City, State Zip (Street Address block is removed)
    if contact_info.get("city_state_zip"):
        add_styled_paragraph(
            document,
            contact_info["city_state_zip"],
            font_name=letterhead_font_name,
            font_size=letterhead_font_size,
            alignment=letterhead_alignment,
            space_after=Pt(1)
        )

    # Phone Number
    phone_top = contact_info.get("phone")
    if phone_top:
        add_styled_paragraph(
            document,
            phone_top,
            font_name=letterhead_font_name,
            font_size=letterhead_font_size,
            alignment=letterhead_alignment,
            space_after=Pt(1)
        )

    # Email Address (Hyperlinked)
    email_top = contact_info.get("email")
    if email_top:
        p_email_top = document.add_paragraph()
        p_email_top.alignment = letterhead_alignment
        p_email_top.paragraph_format.space_before = Pt(0)
        p_email_top.paragraph_format.space_after = Pt(1)
        p_email_top.paragraph_format.widow_control = True
        add_hyperlink(p_email_top, f"mailto:{email_top}", email_top,
                      font_name=letterhead_font_name, font_size=letterhead_font_size,
                      color_hex="0563C1", is_underline=True)

    # LinkedIn Profile (Hyperlinked)
    if contact_info.get("linkedin_url") and contact_info.get("linkedin_text"):
        p_linkedin_top = document.add_paragraph()
        p_linkedin_top.alignment = letterhead_alignment
        p_linkedin_top.paragraph_format.space_before = Pt(0)
        p_linkedin_top.paragraph_format.space_after = Pt(18) # Space after letterhead before main content
        p_linkedin_top.paragraph_format.widow_control = True
        add_hyperlink(p_linkedin_top, contact_info["linkedin_url"], contact_info["linkedin_text"],
                      font_name=letterhead_font_name, font_size=letterhead_font_size,
                      color_hex="0563C1", is_underline=True)
    else: 
        p_spacer = document.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(18) # Ensure consistent spacing if LinkedIn is missing

    # --- Process the main body of the cover letter from LLM ---
    # (This part remains the same: stripping LLM signature, adding paragraphs)
    # ... (body_content_from_llm, closing_pattern, main_body_text, visual_paragraphs loop) ...
    body_content_from_llm = cover_letter_body_text
    candidate_name_from_contact = contact_info.get('name', 'Venkatesh Shanmugam')
    candidate_name_for_regex = re.escape(candidate_name_from_contact)
    closing_pattern = re.compile(rf"(?i)\bSincerely,?\s*(\n\s*)*{candidate_name_for_regex}\s*$", re.MULTILINE)
    main_body_text = closing_pattern.sub("", body_content_from_llm).strip()

    if main_body_text:
        visual_paragraphs = re.split(r'\n{2,}', main_body_text)
        for para_text_segment in visual_paragraphs:
            processed_para_text = para_text_segment.replace('\n', ' ').strip()
            if processed_para_text:
                add_styled_paragraph(document, processed_para_text,
                                     font_name='Times New Roman', font_size=Pt(11),
                                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                     space_after=Pt(8), line_spacing=1.15)
    else:
        add_styled_paragraph(document, "[Cover letter body content was not generated or was stripped with the signature.]",
                             font_name='Times New Roman', font_size=Pt(11))


    # --- Add Controlled Closing, Signature, and Contact Details (as refined previously) ---
    # (This includes: Sincerely, blank line, Name, Phone, Email, GitHub, Portfolio)
    # ... (The detailed closing block code from the previous full function response) ...
    closing_font_name_sig = 'Times New Roman' 
    closing_font_size_sig = Pt(11)

    sincerely_p = document.add_paragraph("Sincerely,")
    sincerely_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sincerely_p.paragraph_format.space_before = Pt(12)
    sincerely_p.paragraph_format.space_after = Pt(0)
    sincerely_p.paragraph_format.widow_control = True
    for run in sincerely_p.runs:
        run.font.name = closing_font_name_sig
        run.font.size = closing_font_size_sig

    signature_space_p = document.add_paragraph()
    run_sig_space = signature_space_p.add_run()
    run_sig_space.font.size = closing_font_size_sig
    signature_space_p.paragraph_format.space_before = Pt(0)
    signature_space_p.paragraph_format.space_after = Pt(0)

    name_p = document.add_paragraph(candidate_name_from_contact)
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_before = Pt(2)
    name_p.paragraph_format.space_after = Pt(0)
    name_p.paragraph_format.widow_control = True
    for run in name_p.runs:
        run.font.name = closing_font_name_sig
        run.font.size = closing_font_size_sig

    phone_number_closing = contact_info.get("phone")
    if phone_number_closing:
        p_phone_closing = document.add_paragraph()
        p_phone_closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_phone_closing.paragraph_format.space_before = Pt(2)
        p_phone_closing.paragraph_format.space_after = Pt(0)
        p_phone_closing.paragraph_format.widow_control = True
        run_phone_cl = p_phone_closing.add_run(phone_number_closing)
        run_phone_cl.font.name = closing_font_name_sig
        run_phone_cl.font.size = closing_font_size_sig

    email_address_closing = contact_info.get("email")
    if email_address_closing:
        p_email_closing = document.add_paragraph()
        p_email_closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_email_closing.paragraph_format.space_before = Pt(2)
        p_email_closing.paragraph_format.space_after = Pt(0)
        p_email_closing.paragraph_format.widow_control = True
        add_hyperlink(p_email_closing, f"mailto:{email_address_closing}", email_address_closing,
                      font_name=closing_font_name_sig, font_size=closing_font_size_sig,
                      color_hex="0563C1", is_underline=True)

    space_before_profile_links = Pt(2)
    if not phone_number_closing and not email_address_closing:
        space_before_profile_links = Pt(6)

    if contact_info.get("github_url") and contact_info.get("github_text"):
        p_github_cl_sig = document.add_paragraph()
        p_github_cl_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_github_cl_sig.paragraph_format.space_before = space_before_profile_links
        p_github_cl_sig.paragraph_format.space_after = Pt(0)
        p_github_cl_sig.paragraph_format.widow_control = True
        add_hyperlink(p_github_cl_sig, contact_info["github_url"], contact_info.get("github_text", "GitHub"),
                      font_name=closing_font_name_sig, font_size=closing_font_size_sig, color_hex="0563C1", is_underline=True)
        space_before_profile_links = Pt(2)

    if contact_info.get("portfolio_url") and contact_info.get("portfolio_text"):
        p_portfolio_cl_sig = document.add_paragraph()
        p_portfolio_cl_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_portfolio_cl_sig.paragraph_format.space_before = space_before_profile_links
        p_portfolio_cl_sig.paragraph_format.space_after = Pt(0)
        p_portfolio_cl_sig.paragraph_format.widow_control = True
        add_hyperlink(p_portfolio_cl_sig, contact_info["portfolio_url"], contact_info.get("portfolio_text", "Portfolio"),
                      font_name=closing_font_name_sig, font_size=closing_font_size_sig, color_hex="0563C1", is_underline=True)

    # --- Construct the base filename for the PDF ---
    candidate_last_name = contact_info.get("name", "Candidate").split()[-1] if contact_info.get("name") else "CL"
    company_str = re.sub(r'\W+', '', company_name) if company_name else "TargetCompany"
    base_pdf_filename = f"{filename_keyword}_{company_str}_{candidate_last_name}"
    base_pdf_filename = re.sub(r'[^\w\.\-_]', '_', base_pdf_filename)

    return generate_pdf_via_google_drive(document, output_pdf_directory, base_pdf_filename)